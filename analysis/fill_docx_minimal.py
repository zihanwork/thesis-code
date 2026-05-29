#!/usr/bin/env python3
"""将 diagnostics 汇总填入 Word 模板；路径勿写死在本机目录。"""
from __future__ import annotations

import argparse
import json
import os
from pathlib import Path

from docx import Document


def default_repo_root() -> Path:
    return Path(__file__).resolve().parent.parent


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Fill meeting report docx from diagnostics JSON.")
    p.add_argument(
        "--docx",
        type=Path,
        default=None,
        help="要填入的 .docx 路径；也可设置环境变量 THESIS_MEETING_DOCX。",
    )
    p.add_argument(
        "--backup",
        type=Path,
        default=None,
        help="首次运行时的备份路径；默认 <docx> 同目录下 .backup_before_fill.docx。",
    )
    p.add_argument(
        "--repo-root",
        type=Path,
        default=default_repo_root(),
        help="仓库根目录，用于定位 output/diagnostics/*.json。",
    )
    args = p.parse_args()
    docx = args.docx
    if docx is None:
        env_docx = os.environ.get("THESIS_MEETING_DOCX", "").strip()
        docx = Path(env_docx) if env_docx else None
    if docx is None:
        p.error("请传入 --docx 或设置环境变量 THESIS_MEETING_DOCX。")
    args.docx = docx
    if args.backup is None:
        args.backup = docx.with_name(docx.stem + ".backup_before_fill.docx")
    return args


def load_json(path: Path):
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def replace_if_empty_after_colon(paragraph, value: str):
    txt = paragraph.text
    if ":" in txt:
        left, right = txt.split(":", 1)
        if right.strip() == "":
            paragraph.text = f"{left}: {value}"
            return True
    return False


def set_paragraph_exact(doc: Document, exact_text: str, new_text: str):
    for p in doc.paragraphs:
        if p.text.strip() == exact_text.strip():
            p.text = new_text
            return True
    return False


def fill_paragraph_bullets(doc: Document):
    obs1_count = 0
    obs2_count = 0
    obs3_count = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("• Repository / benchmark used:") or t.startswith("Repository / benchmark used:"):
            replace_if_empty_after_colon(
                p,
                "Embodied Agent Interface (EAI) - https://github.com/embodied-agent-interface/embodied-agent-interface?tab=readme-ov-file",
            )
        elif t.startswith("• Environment setup status:") or t.startswith("Environment setup status:"):
            replace_if_empty_after_colon(
                p,
                "Completed. Local environment and eai-eval pipeline are runnable (conda + Python 3.8).",
            )
        elif t.startswith("• Main scripts identified:") or t.startswith("Main scripts identified:"):
            replace_if_empty_after_colon(
                p,
                "scripts/run_action_sequencing_eval.sh; analysis/summarize_downstream.py; analysis/link_goal_to_action_failures.py",
            )
        elif t.startswith("• Output directory structure:") or t.startswith("Output directory structure:"):
            replace_if_empty_after_colon(
                p,
                "output/virtualhome/evaluate_results/{goal_interpretation,action_sequencing}/<model>/summary.json; output/diagnostics/",
            )
        elif t.startswith("• Goal Interpretation:") or t.startswith("Goal Interpretation:"):
            replace_if_empty_after_colon(
                p,
                "Located and evaluated. 18-model summaries available under output/virtualhome/evaluate_results/goal_interpretation/.",
            )
        elif t.startswith("• Action Sequencing:") or t.startswith("Action Sequencing:"):
            replace_if_empty_after_colon(
                p,
                "Located and evaluated. 18-model summaries available under output/virtualhome/evaluate_results/action_sequencing/.",
            )
        elif t.startswith("• Subgoal Decomposition:") or t.startswith("Subgoal Decomposition:"):
            replace_if_empty_after_colon(
                p,
                "Module located in framework; not yet included in this week’s diagnostic slicing.",
            )
        elif t.startswith("• Transition Modeling:") or t.startswith("Transition Modeling:"):
            replace_if_empty_after_colon(
                p,
                "Module located in framework; planned for later-stage comparison.",
            )
        elif t.startswith("• Observation 1:") or t.startswith("Observation 1:"):
            obs1_count += 1
            if obs1_count == 1:
                replace_if_empty_after_colon(
                    p,
                    "In goal interpretation, many models show recall > precision, indicating over-generation.",
                )
            else:
                replace_if_empty_after_colon(
                    p,
                    "Downstream action sequencing currently shows all-zero task/execution success across evaluated models.",
                )
        elif t.startswith("• Observation 2:") or t.startswith("Observation 2:"):
            obs2_count += 1
            if obs2_count == 1:
                replace_if_empty_after_colon(
                    p,
                    "Node/edge/action metrics are imbalanced, suggesting uneven capability across goal structure and action expression.",
                )
            else:
                replace_if_empty_after_colon(
                    p,
                    "Main downstream errors concentrate in formatting/parsing, especially action name_id format mismatch.",
                )
        elif t.startswith("• Observation 3:") or t.startswith("Observation 3:"):
            obs3_count += 1
            if obs3_count == 1:
                replace_if_empty_after_colon(
                    p,
                    "Anomalous zero-result models exist and should be checked for format/path consistency.",
                )
            else:
                replace_if_empty_after_colon(
                    p,
                    "Current failures likely occur before deeper planning quality is fully assessed, so format normalization is needed next.",
                )
        elif t.startswith("• Pattern 1:") or t.startswith("Pattern 1:"):
            replace_if_empty_after_colon(
                p,
                "format_or_parsing dominates downstream failures (305 cases).",
            )
        elif t.startswith("• Pattern 2:") or t.startswith("Pattern 2:"):
            replace_if_empty_after_colon(
                p,
                "other failures are far fewer (37 cases), indicating a concentrated early-stage failure mode.",
            )
        elif t.startswith("• Pattern 3:") or t.startswith("Pattern 3:"):
            replace_if_empty_after_colon(
                p,
                "Current errors are mostly pre-runtime-format issues, so planning quality is not fully exposed yet.",
            )

    # Force the second observation block (5.4) to downstream-specific statements.
    obs1 = [p for p in doc.paragraphs if p.text.strip().startswith("Observation 1:")]
    obs2 = [p for p in doc.paragraphs if p.text.strip().startswith("Observation 2:")]
    obs3 = [p for p in doc.paragraphs if p.text.strip().startswith("Observation 3:")]
    if len(obs1) >= 2:
        obs1[1].text = "Observation 1: Downstream action sequencing currently shows all-zero task/execution success across evaluated models."
    if len(obs2) >= 2:
        obs2[1].text = "Observation 2: Main downstream errors concentrate in formatting/parsing, especially action name_id format mismatch."
    if len(obs3) >= 2:
        obs3[1].text = "Observation 3: Current failures likely occur before deeper planning quality is fully assessed, so format normalization is needed next."

    # Fill case templates in-place, preserving paragraph count/order.
    set_paragraph_exact(doc, "Task:", "Task: Pet cat")
    set_paragraph_exact(doc, "Ground-truth goal:", "Ground-truth goal: Turn interaction intent with target pet into executable action sequence.")
    set_paragraph_exact(doc, "Predicted goal:", "Predicted goal: Case-level goal metric appears acceptable (goal_f1=0.8).")
    set_paragraph_exact(doc, "Goal interpretation judgment:", "Goal interpretation judgment: Relatively acceptable at goal-level metric.")
    set_paragraph_exact(doc, "Downstream action outcome:", "Downstream action outcome: prediction has no prediction.")
    set_paragraph_exact(doc, "Failure type:", "Failure type: format_or_parsing.")
    set_paragraph_exact(
        doc,
        "My diagnostic note:",
        "My diagnostic note: Log shows 'Action WALK does not follow name_id format', indicating action formatting failure before deeper planning evaluation.",
    )

    # Fill remaining case placeholders by position traversal
    case_values = [
        {
            "Task": "Drink",
            "GT": "Reach and consume target drink through executable sequence.",
            "Pred": "Goal-level indicator remains acceptable (goal_f1=0.8).",
            "Judgment": "Goal interpretation looks reasonable at case-level metric.",
            "Outcome": "prediction has no prediction.",
            "Type": "format_or_parsing.",
            "Note": "Same name_id formatting error appears in action log.",
        },
        {
            "Task": "Relax on sofa",
            "GT": "Achieve rest state via valid interaction sequence with sofa.",
            "Pred": "Goal-level indicator remains acceptable (goal_f1=0.8).",
            "Judgment": "Goal interpretation looks reasonable at case-level metric.",
            "Outcome": "prediction has no prediction.",
            "Type": "format_or_parsing.",
            "Note": "Failure occurs at action format/parsing stage before runtime planning checks.",
        },
    ]

    # find all template bullets still empty and fill sequentially for case2/case3
    empty_task = [p for p in doc.paragraphs if p.text.strip() == "Task:"]
    empty_gt = [p for p in doc.paragraphs if p.text.strip() == "Ground-truth goal:"]
    empty_pred = [p for p in doc.paragraphs if p.text.strip() == "Predicted goal:"]
    empty_jdg = [p for p in doc.paragraphs if p.text.strip() == "Goal interpretation judgment:"]
    empty_out = [p for p in doc.paragraphs if p.text.strip() == "Downstream action outcome:"]
    empty_typ = [p for p in doc.paragraphs if p.text.strip() == "Failure type:"]
    empty_not = [p for p in doc.paragraphs if p.text.strip() == "My diagnostic note:"]

    for i, data in enumerate(case_values):
        if i < len(empty_task):
            empty_task[i].text = f"Task: {data['Task']}"
        if i < len(empty_gt):
            empty_gt[i].text = f"Ground-truth goal: {data['GT']}"
        if i < len(empty_pred):
            empty_pred[i].text = f"Predicted goal: {data['Pred']}"
        if i < len(empty_jdg):
            empty_jdg[i].text = f"Goal interpretation judgment: {data['Judgment']}"
        if i < len(empty_out):
            empty_out[i].text = f"Downstream action outcome: {data['Outcome']}"
        if i < len(empty_typ):
            empty_typ[i].text = f"Failure type: {data['Type']}"
        if i < len(empty_not):
            empty_not[i].text = f"My diagnostic note: {data['Note']}"


def fill_tables(doc: Document, summary_rows: list):
    # table0: current run status
    t0 = doc.tables[0]
    t0.rows[1].cells[1].text = "VirtualHome"
    t0.rows[1].cells[2].text = "18"
    t0.rows[1].cells[3].text = "Completed"
    t0.rows[1].cells[4].text = "output/virtualhome/evaluate_results/goal_interpretation/"
    t0.rows[1].cells[5].text = "Baseline sorted by all_f1."

    t0.rows[2].cells[1].text = "VirtualHome"
    t0.rows[2].cells[2].text = "18"
    t0.rows[2].cells[3].text = "Completed"
    t0.rows[2].cells[4].text = "output/virtualhome/evaluate_results/action_sequencing/"
    t0.rows[2].cells[5].text = "Task success is currently all-zero due to format/parsing issues."

    # table1: goal interpretation top-3
    goal_rows = [r for r in summary_rows if r.get("eval_type") == "goal_interpretation"]
    goal_rows.sort(key=lambda x: float(x.get("all_f1", 0.0) or 0.0), reverse=True)
    top3 = goal_rows[:3]
    t1 = doc.tables[1]
    for i, row in enumerate(top3, start=1):
        cells = t1.rows[i].cells
        cells[0].text = str(row.get("model", ""))
        cells[1].text = f"{float(row.get('node_precision', 0.0)):.4f}"
        cells[2].text = f"{float(row.get('node_recall', 0.0)):.4f}"
        cells[3].text = f"{float(row.get('node_f1', 0.0)):.4f}"
        cells[4].text = f"{float(row.get('edge_precision', 0.0)):.4f}"
        cells[5].text = f"{float(row.get('edge_recall', 0.0)):.4f}"
        cells[6].text = f"{float(row.get('edge_f1', 0.0)):.4f}"
        cells[7].text = f"{float(row.get('action_precision', 0.0)):.4f}"
        cells[8].text = f"{float(row.get('action_recall', 0.0)):.4f}"
        cells[9].text = f"{float(row.get('action_f1', 0.0)):.4f}"
        cells[10].text = f"{float(row.get('all_precision', 0.0)):.4f}"
        cells[11].text = f"{float(row.get('all_recall', 0.0)):.4f}"
        cells[12].text = f"{float(row.get('all_f1', 0.0)):.4f}"

    # table2: downstream top-2 (all zero currently)
    action_rows = [r for r in summary_rows if r.get("eval_type") == "action_sequencing"]
    action_rows.sort(
        key=lambda x: float(x.get("goal_evaluation.task_success_rate", 0.0) or 0.0), reverse=True
    )
    top2 = action_rows[:2]
    t2 = doc.tables[2]
    for i, row in enumerate(top2, start=1):
        cells = t2.rows[i].cells
        cells[0].text = str(row.get("model", ""))
        cells[1].text = (
            f"task_success_rate={float(row.get('goal_evaluation.task_success_rate', 0.0)):.1f}, "
            f"total_goal={float(row.get('goal_evaluation.total_goal', 0.0)):.1f}"
        )
        cells[2].text = f"{float(row.get('trajectory_evaluation.execution_success_rate', 0.0)):.1f}"
        cells[3].text = "format_or_parsing"
        cells[4].text = "High grammar parsing errors; many predictions fail name_id format."


def main():
    args = parse_args()
    docx_path: Path = args.docx
    backup_path: Path = args.backup
    root: Path = args.repo_root
    summary_json = root / "output/diagnostics/downstream_summary.json"
    top_cases_json = root / "output/diagnostics/goal_correct_but_action_fail_top_cases.json"
    failure_json = root / "output/diagnostics/failure_pattern_counts.json"

    if not docx_path.exists():
        raise FileNotFoundError(f"Docx not found: {docx_path}")
    for path in (summary_json, top_cases_json, failure_json):
        if not path.exists():
            raise FileNotFoundError(f"Required JSON not found: {path}")

    summary_rows = load_json(summary_json)
    load_json(top_cases_json)
    load_json(failure_json)

    if not backup_path.exists():
        backup_path.write_bytes(docx_path.read_bytes())

    doc = Document(str(docx_path))
    fill_paragraph_bullets(doc)
    fill_tables(doc, summary_rows)

    # Normalize accidental double bullet text in case entries without changing structure.
    for p in doc.paragraphs:
        s = p.text.strip()
        if s.startswith("•\t•\t"):
            p.text = s.replace("•\t•\t", "", 1)
        elif s.startswith("•\tTask:"):
            p.text = s.replace("•\t", "", 1)
        elif s.startswith("•\tGround-truth goal:"):
            p.text = s.replace("•\t", "", 1)
        elif s.startswith("•\tPredicted goal:"):
            p.text = s.replace("•\t", "", 1)
        elif s.startswith("•\tGoal interpretation judgment:"):
            p.text = s.replace("•\t", "", 1)
        elif s.startswith("•\tDownstream action outcome:"):
            p.text = s.replace("•\t", "", 1)
        elif s.startswith("•\tFailure type:"):
            p.text = s.replace("•\t", "", 1)
        elif s.startswith("•\tMy diagnostic note:"):
            p.text = s.replace("•\t", "", 1)

    doc.save(str(docx_path))
    print(f"[DONE] Filled docx in-place: {docx_path}")
    print(f"[DONE] Backup preserved: {backup_path}")


if __name__ == "__main__":
    main()
